import os
import pandas as pd
import streamlit as st
from moviepy.editor import *
import whisper
import torch
import librosa
import google.generativeai as genai
import boto3
from botocore.exceptions import NoCredentialsError
from datetime import datetime
from docx import Document
from io import BytesIO
import markdown
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from tempfile import NamedTemporaryFile
import tempfile
from pydub import AudioSegment
# import openai
# import openai_whisper

# AWS S3 configuration
AWS_ACCESS_KEY = ''
AWS_SECRET_KEY = ''
BUCKET_NAME_INP = 'aimedscribe'
BUCKET_NAME_OUT = 'aimedscribe'
input_file_path = ''



# Function to convert mp4 to mp3
def convert_to_mp3(mp4_file,filepath):
    # Convert mp4 to mp3 (you need to implement this)
    print("audio file generation in progress")
    print(mp4_file.name)
    print(filepath)
    vidfile=VideoFileClip(mp4_file.name)
    vidfile.audio.write_audiofile(filepath)
    #vidfile.audio.write_audiofile(filepath+"\\audfile.mp3")
    print("audio file generated")

# Function to generate transcript using Whisper API
def generate_transcript(audio_file):
    print("transcript generated started")
    print(audio_file)
    
    audfile,sr = librosa.load(audio_file, sr=16000)
    device="cuda:0" if torch.cuda.is_available() else "cpu"
    print(device)
    print(torch.cuda.is_available())
    model=whisper.load_model("base.en",device=device)
    transcript = model.transcribe(audio=audfile)
    print("transcript generated")
    return transcript

#Function to upload input file to S3
def upload_to_s3(file, object_name=None):
    # If S3 object_name was not specified, use file name
    if object_name is None:
        object_name = file.name

    # Initialize a session using Amazon S3
    s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                             aws_secret_access_key=AWS_SECRET_KEY)

    try:
        # Upload the file
        print("file",file)
        print("bucket name",BUCKET_NAME_INP)
        print("obj",object_name)
        s3_client.upload_fileobj(file, BUCKET_NAME_INP, object_name)
        st.success(f"File {file.name} uploaded to {BUCKET_NAME_INP}/{object_name} successfully!")
        global input_file_path
        input_file_path = BUCKET_NAME_INP + "/" + object_name
        input_file_path = object_name
        print(input_file_path,"input_file_path")
        print("upload success")
    except NoCredentialsError:
        st.error("Credentials not available")

def upload_file_to_s31(file_name, object_name=None):
    """
    Uploads a file to an S3 bucket.

    :param file_name: File to upload
    :param bucket_name: Name of the S3 bucket
    :param object_name: S3 object name. If not specified, file_name is used.
    :return: True if file was uploaded, else False
    """
    # Use the object_name if specified; otherwise, default to file_name
    if object_name is None:
        object_name = file_name.name

    print("file",file_name)
    print("bucket name",BUCKET_NAME_INP)
    print("obj",object_name)

    # Create an S3 client
    s3_client = boto3.client('s3')

    try:
        # Upload the file to S3
        # response = s3_client.upload_file(file_name, BUCKET_NAME_INP, object_name)
        file_name.seek(0)
        s3_client.upload_fileobj(file_name, BUCKET_NAME_INP, object_name)
        print(f"File '{file_name}' uploaded to S3 bucket '{BUCKET_NAME_INP}' as '{object_name}'.")
        global input_file_path
        input_file_path = BUCKET_NAME_INP + "/" + object_name
        input_file_path = object_name
        print(input_file_path,"input_file_path")
        print("upload success")
        return True
    except FileNotFoundError:
        print(f"The file '{file_name}' was not found.")
        return False
    # except NoCredentialsError:
    #     print("AWS credentials not available.")
    #     return False
    # except PartialCredentialsError:
    #     print("Incomplete AWS credentials provided.")
    #     return False
    except Exception as e:
        print(f"An error occurred: {e}")
        return False

# #####################

def pull_files_s3(key) -> str:
    try:
        # st.write("starrt")
        # st.write(BUCKET_NAME_INP)
        # st.write(key)
        s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                             aws_secret_access_key=AWS_SECRET_KEY)
        print(key,"key")
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            # st.write("inside download from s3")
            s3_client.download_fileobj(BUCKET_NAME_INP, key, temp_file)
            temp_file.flush()
            temp_file.seek(0)
            st.write("File downloaded successfully!")
            return temp_file.name
    except Exception as e:
        # st.error("Unable to download file from s3") 
        st.error(e)
    
# function to extract audio from video
def extract_audio(key) -> str:
    
    temp_vdo_path = None
    try:
        # st.write("Inside Audio")
        # st.write(key)
        temp_vdo_path = pull_files_s3(key)
        # st.write("here we go with the path")
        video = VideoFileClip(temp_vdo_path)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
            video.audio.write_audiofile(temp_audio.name)
            temp_audio_path = temp_audio.name
        video.close()
        st.write("mp3 generated")    
       
        
    except Exception as e:
        st.error("Unable to convert to mp3")    
        
    finally:
        if temp_vdo_path and os.path.exists(temp_vdo_path):
            os.remove(temp_vdo_path)                      
    return temp_audio_path




def transcribe_audio(key) -> str:
   
    try:
        if key.lower().startswith('input'):
            temp_audio_path = pull_files_s3(key)
        else:
            temp_audio_path = key

        # st.write(temp_audio_path)

        audfile,sr = librosa.load(temp_audio_path, sr=16000)
        device="cuda:0" if torch.cuda.is_available() else "cpu"
        # st.write(device)
        # st.write(torch.cuda.is_available())
        # st.write("start ttranscript gen")
        model=whisper.load_model("base.en",device=device)
        transcript = model.transcribe(audio=audfile)
        st.write("transcript generated")
        # return transcript
        
    except Exception as e:
        st.error(e)
        st.error("transcript generation failed")
        
    finally:  
        if temp_audio_path and os.path.exists(temp_audio_path):
            os.remove(temp_audio_path)
        st.write("Temporary audio file removed successfully!")            
        return transcript


# ##############

def pull_files_s31(key) -> str:
    try:
        # Extract the file extension from the S3 key
        file_extension = key.split('.')[-1]
        if not file_extension:
            raise ValueError("Key does not have a file extension!")

        # Create a temporary file with the same extension
        with tempfile.NamedTemporaryFile(suffix=f".{file_extension}", delete=False) as temp_file:
            s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                                     aws_secret_access_key=AWS_SECRET_KEY)
            s3_client.download_fileobj(BUCKET_NAME_INP, key, temp_file)
            temp_file.flush()
            temp_file.seek(0)
            st.write("File downloaded successfully!")
            return temp_file.name
    except Exception as e:
        st.error("Unable to download file from S3")
        st.error(e)




# Function to upload output to S3 
def upload_string_to_s3(content, object_name):
    # Initialize a session using Amazon S3
    s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                             aws_secret_access_key=AWS_SECRET_KEY)

    try:
        # Upload the string content
        s3_client.put_object(Bucket=BUCKET_NAME_OUT, Key=object_name, Body=content)
        st.success(f"Content uploaded to {BUCKET_NAME_OUT}/{object_name} successfully!")       
    except NoCredentialsError:
        st.error("Credentials not available")

def upload_to_s31(file_path, object_name):
    s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                             aws_secret_access_key=AWS_SECRET_KEY)
    try:
        s3_client.upload_file(file_path, BUCKET_NAME_OUT, object_name)
        st.success(f"File uploaded to S3 bucket '{BUCKET_NAME_OUT}'/'{object_name}'")
    except Exception as e:
        st.error(f"Failed to upload file to S3: {e}")

# Function to add content from HTML to Word Document
def add_html_to_doc(element, doc):
    if element.name == 'h1':
        paragraph = doc.add_heading(level=1)
        paragraph.text = element.get_text()
    elif element.name == 'h2':
        paragraph = doc.add_heading(level=2)
        paragraph.text = element.get_text()
    elif element.name == 'h3':
        paragraph = doc.add_heading(level=3)
        paragraph.text = element.get_text()
    elif element.name == 'p':
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(element.get_text())
        # Check for strong (bold) text
        if element.find('strong'):
            run.bold = True
    elif element.name == 'strong':
        run = doc.add_paragraph().add_run(element.get_text())
        run.bold = True

        
#Function to upload output to S3 in docx
def string_to_word_and_upload(content, object_name):
    # Create a new Document
    #html_content = markdown.markdown(content)
    html_content=content

    # Parse HTML Content
   # soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Document
    doc = Document()

    # Add HTML content to the Document
    doc.add_paragraph(html_content)

    #for element in soup.body:
     #   add_html_to_doc(element, doc)

    # Save the Document to a BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # Upload the Word document to S3
    upload_string_to_s3(byte_io, object_name)


#Function to upload output to S3 in docx
def string_to_word_and_upload1(content, object_name):

    document_path = create_word_document(content)
    upload_to_s31(document_path, object_name)
    with open(document_path, "rb") as doc_file:
        btn = st.download_button(
                        label="Download",
                        data=doc_file,
                        file_name="res.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        
        doc_file.close()

# Function to generete content using Gemini
def generate_content(transcript):
     
        
     print("generation started")
     prompt = f"""
     Imagine you are content generator to update the medical summary of the patient to the databse and your job is to read 
     the transcript and generate the content

     Capturing the details accurately is very critical as this will be used by the doctors in future 
     to understand the current status and next course of action.

     The response should have the following sections:
     1. Chief Complaint/Reason for Visit  
     2. Presenting Symptoms  
     3. Relevant Medical History  
     4. Current Medications  
     5. Allergies  
     6. Vital Signs  
     7. Physical Examination Findings  
     8. Laboratory and Diagnostic Results  
     9. Diagnosis  
     10. Treatment Plan  
     11. Medication Changes  
     12. Surgical Procedures Planned/Performed  
     13. Intraoperative Findings (if surgery performed)  
     14. Post-Surgical Status and Care  
     15. Follow-up Plans and Recommendations  
     16. Patient Education and Counseling  
     17. Discharge Instructions (if applicable)  
     18. Referrals and Consultations
     
     Generate content based on the following transcript 

     {transcript}

     """
   
     # prompt=prompt1

     genai.configure(api_key='AIzaSyDJZkM2KyZT_wQIbm2ruDHuOxPPXvrAqwA')

     # Set up the model
     generation_config = {
     "temperature":0,
     "top_p": 0.95,
     "top_k": 0,
     "max_output_tokens": 30000,
     }

     safety_settings = [
     {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
     },
     {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
     },
     {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
     },
     {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
     },
     ]

      
     model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest",
                              generation_config=generation_config,
                              safety_settings=safety_settings)

     convo = model.start_chat(history=[
     ])

     convo.send_message(prompt)
     reslt = convo.last.text
     print('generation completed')
     return reslt


# function to write it into a wordfile
def create_word_document(text):
    doc = Document()
    doc.add_paragraph(text)
    with NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        doc.save(temp_file.name)
    return temp_file.name


import os
import tempfile
import boto3

def pull_files_s32(key) -> str:
    try:
        file_extension = key.split('.')[-1]
        if not file_extension:
            raise ValueError("Key does not have a file extension!")

        # Create a persistent temporary file
        with tempfile.NamedTemporaryFile(suffix=f".{file_extension}", delete=False) as temp_file:
            temp_file_name = temp_file.name

        s3_client = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY,
                                 aws_secret_access_key=AWS_SECRET_KEY)
        with open(temp_file_name, 'wb') as f:
            s3_client.download_fileobj(BUCKET_NAME_INP, key, f)

        # Validate file existence and size
        if not os.path.exists(temp_file_name):
            raise FileNotFoundError(f"File {temp_file_name} does not exist!")
        if os.path.getsize(temp_file_name) == 0:
            raise ValueError(f"File {temp_file_name} is empty!")

        return temp_file_name

    except Exception as e:
        raise RuntimeError(f"Unable to download or process file: {e}")



def main():
    st.title(r"AI Medical Scribe")
    
    # File upload section
    st.sidebar.title("Upload File")
    
    #sess = sagemaker.Session()
    #bucket = sess.default_bucket() # Set a default S3 bucket
    #prefix = '/input/'

    #st.write('s3')
    uploaded_file = st.sidebar.file_uploader("Upload Recording", type=["mp4", "mp3","txt"])
    #st.write('s31')
    # global temp_audio_path
    # Button to trigger generation
    if st.sidebar.button("Generate"):


        
        if uploaded_file is not None:
            if uploaded_file.name.endswith('.txt'):
                uploaded_file1=uploaded_file.read()
            if uploaded_file.name.endswith('.mp3'):
                uploaded_file1=uploaded_file.read()
            obj_name_inp = f"input/{uploaded_file.name}"
            


            # Generate a timestamp
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            # Construct the object name with timestamp
            file_extension = uploaded_file.name.split('.')[-1]
            #obj_name_inp = f"input/{uploaded_file.name.split('.')[0]}_{timestamp}.{file_extension}"

            # upload_to_s3(uploaded_file,obj_name_inp)
            print("uploaded_file",uploaded_file)
            print("obj",obj_name_inp)
            print("start")
            upload_file_to_s31(uploaded_file,obj_name_inp)

            if uploaded_file.name.endswith('.txt'):
               transcript = uploaded_file1.decode("utf-8")            
               #with open(uploaded_file.name,'r',encoding='utf-8') as text:
                #    transcript=text.read()
               # st.write('started')
                #with open('/tmp/tempfile.txt', 'wb') as f:
                #    f.write(uploaded_file.getbuffer())
       
                # Upload the file to S3
                #uploaded = upload_to_aws('/tmp/tempfile.txt', bucket , 'transcript.txt')

                #if uploaded:
                #    st.success("File uploaded to S3")
              #  else:
                 #   st.error("Failed to upload the file to S3")
            else:
                if uploaded_file.name.endswith('.mp4'):                    
                    st.write("Converting mp4 to mp3...")
                    input_file_path_audio = input_file_path.replace('.mp4','.mp3')
                    filepath="audfile.mp3"
                    # st.write(uploaded_file)
                    # st.write("inp")
                    # st.write(input_file_path)
                    # st.write("actual call")
                    
                    # audio_file = convert_to_mp3(uploaded_file,filepath)
                    

                    file_key = input_file_path

                    print("1",file_key)
                    # global temp_audio_path
                    temp_audio_path = extract_audio(file_key)
                    print("2",temp_audio_path)
                    # st.write(temp_audio_path)
                    
                    
                    # st.write("audio file generated")
                    
                    # transcript = transcribe_audio(temp_audio_path)
                    

                
                

                




                
                elif uploaded_file.name.endswith('.mp3'):
                    #audio_file = uploaded_file
                    #code needs to be corrected
                    file_key = input_file_path
                    st.write(file_key)
                    print("3",file_key)
                    temp_audio_path=pull_files_s31(file_key)
                    print("4",temp_audio_path)
                    # temp_audio_path = temp_audio_path+".mp3"
                    # print("5",temp_audio_path)
                    print(f"File exists: {os.path.exists(temp_audio_path)}")
                    # info = mediainfo(temp_audio_path)
                    # print("info",info)
                    st.write(temp_audio_path)
                    file_size=os.path.getsize(temp_audio_path)
                    print("file_size",file_size)
                    # import soundfile as sf
                    # with sf.SoundFile(temp_audio_path) as audio_file:
                    #     print(f"Audio file is valid: {audio_file.samplerate} Hz")

                    # clean_audio_path = "/tmp/clean_audio.mp3"
                    # print('start1')
                    # audio = AudioSegment.from_file(temp_audio_path, format="mp3")
                    # print('start12')
                    # audio.export(clean_audio_path, format="mp3")
                    # print('start13')
                    # temp_audio_path=clean_audio_path


                    
                    # audfile,sr = librosa.load(temp_audio_path, sr=16000)
                    # device="cuda:0" if torch.cuda.is_available() else "cpu"
                    # print(device)
                    # print(torch.cuda.is_available())
                    # model=whisper.load_model("base.en",device=device)
                    # transcript = model.transcribe(audio=audfile)

                
                    # with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio_file1:
                    #     temp_audio_file1.write(uploaded_file1)
                    #     temp_audio_path = temp_audio_file1.name
                    # # audio_file=input_file_path
                    # temp_audio_path = audio_file


                               
                
                                
                st.write("Generating transcript...")
                # global temp_audio_path
                print(temp_audio_path)
                st.write("Transcript generated")
                transcript = generate_transcript(temp_audio_path)
                st.write("Transcript generated")
                
                # if os.path.exists(temp_audio_path) and os.path.getsize(temp_audio_path) > 0:
                #     st.write("Generating transcript...")
                #     transcript = generate_transcript(temp_audio_path)
                #     st.write("Transcript generated")
                # else:
                #     st.write("Empty File")
            
            
            
            st.write("Generating content...")
            result = generate_content(transcript)
            
            st.write("Content generated:")
                     
            if result:
                
                obj_name_out = f"output/res.txt"                   
                obj_name_out_docx = f"output/res.docx" 
                print(result)
                st.write(result)
                string_to_word_and_upload1(result,obj_name_out_docx)
                st.write("uploaded result to S3")
                

     


if __name__== '__main__':    
    main()
